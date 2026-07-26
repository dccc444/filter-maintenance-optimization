from __future__ import annotations

import re
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / ".tmp_final_paper" / "pydeps"))

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Pt, RGBColor
from latex2mathml.converter import convert as latex_to_mathml
from lxml import etree


OUT = ROOT / "paper" / "过滤设备透水率退化建模与维护策略优化.docx"
COMBINED = ROOT / "paper" / "最终论文正文.md"
PREFACE = ROOT / "paper" / "最终论文前置与附录.md"
CHAPTERS = [
    ROOT / "paper" / "第一问论文正文.md",
    ROOT / "paper" / "第二问论文正文.md",
    ROOT / "paper" / "第三问论文正文.md",
    ROOT / "paper" / "第四问论文正文.md",
]
MML_XSL = Path(r"C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL")


def combine_source() -> str:
    frame = PREFACE.read_text(encoding="utf-8")
    chapters = []
    for path in CHAPTERS:
        text = path.read_text(encoding="utf-8")
        text = re.sub(r"^>.*(?:\n|$)", "", text, flags=re.M)
        chapters.append(text.strip())
    joined = frame.replace(
        "<!-- 四问正文由构建脚本插入此处 -->",
        "\n\n".join(chapters),
    )
    COMBINED.write_text(joined, encoding="utf-8")
    return joined


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_margins(cell, top=55, start=70, bottom=55, end=70):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        run._r.append(node)


def set_run_font(run, cn="宋体", en="Times New Roman", size=10.5, bold=None):
    run.font.name = en
    run._element.rPr.rFonts.set(qn("w:eastAsia"), cn)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def add_inline_runs(paragraph, text: str, bold_all=False):
    parts = re.split(r"(\*\*.*?\*\*|`.*?`|\\\(.*?\\\))", text)
    for part in parts:
        if not part:
            continue
        bold = bold_all
        if part.startswith("**") and part.endswith("**"):
            part = part[2:-2]
            bold = True
        elif part.startswith("`") and part.endswith("`"):
            part = part[1:-1]
        elif part.startswith(r"\(") and part.endswith(r"\)"):
            part = part[2:-2]
        run = paragraph.add_run(part)
        set_run_font(run, cn="宋体", en="Times New Roman", size=10.5, bold=bold)


def add_equation(paragraph, latex: str):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cleaned = latex.strip()
    cleaned = re.sub(r"\\tag\{([^}]+)\}", "", cleaned)
    try:
        mathml = latex_to_mathml(cleaned)
        transform = etree.XSLT(etree.parse(str(MML_XSL)))
        omml = transform(etree.fromstring(mathml.encode("utf-8")))
        paragraph._p.append(parse_xml(etree.tostring(omml)))
    except Exception:
        run = paragraph.add_run(cleaned)
        set_run_font(run, cn="Cambria Math", en="Cambria Math", size=10.5)


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)
    add_page_number(section.footer.paragraphs[0])

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.15
    pf.space_after = Pt(0)
    pf.first_line_indent = Cm(0.74)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for name, size, cn, before, after in [
        ("Title", 22, "黑体", 0, 20),
        ("Heading 1", 16, "黑体", 14, 8),
        ("Heading 2", 14, "黑体", 10, 5),
        ("Heading 3", 12, "黑体", 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), cn)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    doc.styles["Title"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.styles["Heading 1"].paragraph_format.page_break_before = False

    if "CaptionCN" not in [s.name for s in doc.styles]:
        style = doc.styles.add_style("CaptionCN", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.size = Pt(9)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style.paragraph_format.space_before = Pt(3)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True

    settings = doc.settings._element
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    settings.append(update)


def add_table(doc: Document, lines: list[str]):
    rows = []
    for line in lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in rows[1]):
        rows.pop(1)
    table = doc.add_table(rows=len(rows), cols=max(map(len, rows)))
    table.style = "Table Grid"
    table.autofit = False
    width = Cm(16.0) / max(map(len, rows))
    for r_idx, values in enumerate(rows):
        row = table.rows[r_idx]
        if r_idx == 0:
            set_repeat_table_header(row)
        for c_idx, value in enumerate(values):
            cell = row.cells[c_idx]
            cell.width = width
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if r_idx == 0:
                set_cell_shading(cell, "E7EEF6")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = None
            add_inline_runs(p, value, bold_all=(r_idx == 0))
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_image(doc: Document, alt: str, target: str, fig_no: int):
    target = target.replace("?", "")
    image_path = (COMBINED.parent / target).resolve()
    if not image_path.exists():
        candidates = list(image_path.parent.glob(image_path.stem.split(".")[0] + "*"))
        if candidates:
            image_path = candidates[0]
    if image_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.keep_with_next = True
        p.add_run().add_picture(str(image_path), width=Cm(14.8))
        cap = doc.add_paragraph(style="CaptionCN")
        label = re.sub(r"^图\s*\d+\s*", "", alt).strip()
        cap.add_run(f"图 {fig_no}  {label}")


def build(md: str):
    doc = Document()
    configure_document(doc)
    lines = md.splitlines()
    i = 0
    equation = []
    in_equation = False
    figure_count = 0
    image_budget = 12
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        if in_equation:
            if line == r"\]":
                p = doc.add_paragraph()
                p.paragraph_format.keep_together = True
                add_equation(p, "\n".join(equation))
                equation = []
                in_equation = False
            else:
                equation.append(raw)
            i += 1
            continue
        if line == r"\[":
            in_equation = True
            i += 1
            continue
        if not line:
            i += 1
            continue
        if line == r"\newpage":
            doc.add_page_break()
            i += 1
            continue
        if line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_table(doc, table_lines)
            continue
        image = re.fullmatch(r"!\[(.*?)\]\((.*?)\)", line)
        if image:
            figure_count += 1
            if figure_count <= image_budget:
                add_image(doc, image.group(1), image.group(2), figure_count)
            i += 1
            continue
        heading = re.match(r"^(#{1,3})\s+(.*)$", line)
        if heading:
            level = len(heading.group(1))
            text = heading.group(2)
            if level == 1 and not any(p.text for p in doc.paragraphs):
                p = doc.add_paragraph(style="Title")
                add_inline_runs(p, text, bold_all=True)
            else:
                p = doc.add_paragraph(style=f"Heading {level}")
                p.paragraph_format.first_line_indent = None
                add_inline_runs(p, text, bold_all=True)
            i += 1
            continue
        if re.match(r"^\d+\.\s+", line):
            p = doc.add_paragraph(style="Normal")
            p.paragraph_format.left_indent = Cm(0.74)
            p.paragraph_format.first_line_indent = Cm(-0.74)
            add_inline_runs(p, line)
            i += 1
            continue
        if line.startswith("**表") and line.endswith("**"):
            p = doc.add_paragraph(style="CaptionCN")
            p.paragraph_format.keep_with_next = True
            add_inline_runs(p, line[2:-2], bold_all=True)
            i += 1
            continue
        p = doc.add_paragraph(style="Normal")
        add_inline_runs(p, line)
        i += 1

    core = doc.core_properties
    core.title = "过滤设备透水率退化建模与维护策略优化"
    core.subject = "全国大学生数学建模竞赛论文"
    core.author = ""
    core.keywords = "过滤设备, 透水率, 随机退化, 维护优化, 稳健性"
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build(combine_source())
