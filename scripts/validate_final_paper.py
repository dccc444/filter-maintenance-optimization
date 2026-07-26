from pathlib import Path
from zipfile import ZipFile

from docx import Document
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "paper" / "过滤设备透水率退化建模与维护策略优化.docx"
PDF = ROOT / "output" / "pdf" / "过滤设备透水率退化建模与维护策略优化.pdf"

doc = Document(DOCX)
with ZipFile(DOCX) as archive:
    assert archive.testzip() is None

raw_math = [p.text for p in doc.paragraphs if r"\[" in p.text or r"\]" in p.text]
assert not raw_math, f"Unparsed display math remains: {len(raw_math)}"
assert not any("<!--" in p.text for p in doc.paragraphs)

reader = PdfReader(PDF)
assert len(reader.pages) == 23
assert PDF.stat().st_size < 20 * 1024 * 1024
assert DOCX.stat().st_size < 20 * 1024 * 1024
assert (reader.metadata.author or "") == ""

first_page = reader.pages[0].extract_text() or ""
assert "摘要" in first_page and "关键词" in first_page
assert "承诺书" not in first_page and "编号专用页" not in first_page

print(
    {
        "docx_paragraphs": len(doc.paragraphs),
        "docx_tables": len(doc.tables),
        "docx_inline_shapes": len(doc.inline_shapes),
        "pdf_pages": len(reader.pages),
        "docx_bytes": DOCX.stat().st_size,
        "pdf_bytes": PDF.stat().st_size,
        "first_page_is_abstract": True,
    }
)
